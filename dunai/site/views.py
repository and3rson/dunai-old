import os
from pathlib import Path
from io import BytesIO

import yaml
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pylint: disable=no-name-in-module
from docx.enum.section import WD_SECTION  # pylint: disable=no-name-in-module
from docx.enum.style import WD_STYLE_TYPE  # pylint: disable=no-name-in-module
from docx.shared import Pt, RGBColor
from django.urls import reverse
from django.shortcuts import render
from django.conf import settings
from django.http import HttpResponse

from .models import Feature, Project


def get_cv():
    """
    Read & return CV dict from YAML file.
    """
    cvfile = open(os.path.join(settings.BASE_DIR, 'dunai', 'files', 'cv.yaml'))
    with cvfile:
        return yaml.load(cvfile.read())


def index(request):
    """
    Render index page.
    """
    # pylint: disable=no-member
    return render(request, 'index.html', dict(
        cv=get_cv(),
        features=Feature.objects.all(),
        projects=Project.objects.all()
    ))


def cv_doc(request):
    """
    Generate CV in .DOCX format.
    """
    # pylint: disable=invalid-name

    if request.GET.get('preview'):
        url = request.build_absolute_uri(reverse('cv_doc'))
        return HttpResponse(
            '<html>'
            '<body style="margin: 0; padding: 0">'
            '<iframe src="{}" style="width: 100%; height: 100%; '
            'box-sizing: border-box; border: 0 none"></iframe>'
            '</body>'
            '</html>'.format(
                'http://docs.google.com/gview?url={}&embedded=true'.format(url)
            ), content_type='text/html'
        )

    cv = get_cv()

    doc = docx.Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(32)
    date_style = doc.styles.add_style('Date', WD_STYLE_TYPE.CHARACTER)
    date_style.font.color.rgb = RGBColor(192, 192, 216)
    text_par = doc.styles.add_style('Text paragraph', WD_STYLE_TYPE.PARAGRAPH)
    text_par.font.name = 'Arial'
    text_par.paragraph_format.space_before = Pt(8)
    text_par.paragraph_format.first_line_indent = Pt(16)
    h1 = doc.add_heading('Andrew Dunai', 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading('Full Stack + Software Architect', 1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('About', 1)
    for par_text in cv['bio']:
        par = doc.add_paragraph(par_text.strip())
        par.style = text_par

    doc.add_heading('Experience', 1)
    for company in cv['companies']:
        doc.add_section(WD_SECTION.CONTINUOUS)
        head = doc.add_heading('', 2)
        head.add_run(company['role'] + '\n')
        run = head.add_run()
        run.add_picture(str(
            Path(settings.BASE_DIR)
            / 'dunai'
            / 'site'
            / 'static'
            / company['icon']
        ), width=Pt(16), height=Pt(16))
        head.add_run(company['name'])
        run = head.add_run(' ({} - {})'.format(
            company['start'].strftime('%b %Y'),
            company['end'].strftime('%b %Y') if 'end' in company else 'now'
        ))
        run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run.style = date_style
        run.italic = True

        desc = doc.add_paragraph('')
        desc.add_run('About: ').bold = True
        desc.add_run(company['info'].strip())
        desc.style = text_par

        desc = doc.add_paragraph('')
        desc.add_run('My role: ').bold = True
        desc.add_run(company['description'].strip())
        desc.style = text_par

        tech = doc.add_paragraph('')
        tech.style = text_par
        tech.add_run('Technologies: ').bold = True
        for i, technology in enumerate(company['technologies']):
            if i > 0:
                tech.add_run(', ')
            tech.add_run(list(technology.keys())[0])

    doc.add_heading('Fun facts', 1)
    for fact in cv['misc']:
        doc.add_paragraph(fact, style='ListBullet')

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=cv.docx'
    return response
